"""
İhale Dosyası Word Belge Oluşturucu.

Bu modül, kullanıcının form üzerinden girdiği bilgileri kullanarak
orijinal ihale dosyası şablonundaki sarı vurgulu alanları doldurur
ve yeni Word (.docx) ve PDF dosyaları oluşturur.

Temel İşlev:
    - Şablon .docx dosyasını okur
    - Sarı vurgulu run'ları tespit eder
    - Kullanıcı verilerini ilgili alanlara yerleştirir
    - Sarı vurguyu kaldırır (doldurulan alanlardan)
    - Yeni Word ve PDF dosyaları oluşturur
"""

import os
import sys
import re
import subprocess
import platform
from copy import deepcopy
from datetime import datetime, date, time
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# Modül yolunu ayarla (hem local hem cloud ortamında çalışması için)
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from field_config import UNIQUE_FIELDS, YELLOW_TO_UNIQUE_MAP, INSTRUCTION_FIELDS


def get_locative_suffix(word):
    """
    Kelimedeki son ünlüye ve ünsüze göre 'de/'da, 'te/'ta ekini döndürür.
    """
    if not word:
        return ""
    
    # Türkçe karakterleri destekle
    word_lower = word.strip().lower()
    
    # Son ünlü harfi bul
    vowels = "aeıioöuü"
    last_vowel = ""
    for char in reversed(word_lower):
        if char in vowels:
            last_vowel = char
            break
            
    if not last_vowel:
        return "'de" # varsayılan
        
    hard_consonants = "fstkçşhp"
    last_char = word_lower[-1]
    is_hard = last_char in hard_consonants
    
    if last_vowel in "aıou":
        return "'ta" if is_hard else "'da"
    else:
        return "'te" if is_hard else "'de"

def get_dative_suffix(time_str):
    """
    14:00 (on dört'e), 10:30 (otuz'a) vb. için yönelme (ismin -e hali) eki.
    """
    if not time_str:
        return "'e"
    
    # Saat 00 ile bitiyorsa saate bak, yoksa dakikaya bak
    parts = str(time_str).split(":")
    if len(parts) >= 2:
        hour = int(parts[0])
        minute = int(parts[1])
        if minute == 0:
            val = hour
        else:
            val = minute
            
        # Son rakamın okunuşuna göre ek
        last_digit = val % 10
        if val == 0: # sıfır -> a
            return "'a"
        if last_digit in [1, 2, 5, 7, 8]: # bir'e, iki'ye, beş'e, yedi'ye, sekiz'e -> e
            if last_digit == 2 or last_digit == 7:
                return "'ye"
            return "'e"
        elif last_digit in [3, 4, 0]: # üç'e, dört'e, (on'a, otuz'a, kırk'a, altmış'a) -> a/e
            if last_digit == 3 or last_digit == 4:
                return "'e"
            
            # 10, 30, 40, 50, 20
            if val == 10: return "'a" # on'a
            if val == 20: return "'ye" # yirmi'ye
            if val == 30: return "'a" # otuz'a
            if val == 40: return "'a" # kırk'a
            if val == 50: return "'ye" # elli'ye
            
            return "'a"
        elif last_digit in [6, 9]: # altı'ya, dokuz'a
            if last_digit == 6:
                return "'ya"
            return "'a"
            
    return "'e" # varsayılan


def _normalize_text(text):
    """Metni karşılaştırma için normalize eder."""
    if not text:
        return ""
    return text.strip().replace("\u00a0", " ").replace("  ", " ")


def _match_yellow_to_field(yellow_text):
    """
    Sarı vurgulu metni benzersiz alan ID'sine eşleştirir.
    """
    normalized = _normalize_text(yellow_text)

    # Direkt eşleşme
    if normalized in YELLOW_TO_UNIQUE_MAP:
        return YELLOW_TO_UNIQUE_MAP[normalized]

    # Talimat mı kontrol et (Kısmi eşleşmeden ÖNCE)
    for instruction in INSTRUCTION_FIELDS:
        if instruction in normalized or normalized in instruction:
            return None  # Talimat metni, form alanı değil

    # Kısmi eşleşme (sarı metin bir map anahtarını içeriyorsa)
    for key, field_id in YELLOW_TO_UNIQUE_MAP.items():
        if key in normalized or normalized in key:
            return field_id

    return None


def _format_field_value(field_id, value):
    """
    Alan değerini Word'de gösterilecek formata dönüştürür.
    """
    if value is None or value == "":
        return ""

    field_info = UNIQUE_FIELDS.get(field_id, {})
    field_type = field_info.get("type", "text")

    formatted = str(value)
    
    if field_type == "date":
        if isinstance(value, date):
            formatted = value.strftime("%d/%m/%Y")
        elif isinstance(value, str):
            formatted = value
    elif field_type == "time":
        if isinstance(value, time):
            formatted = value.strftime("%H:%M")
        elif isinstance(value, str):
            formatted = value
    elif field_type == "number":
        formatted = str(value)
        
    # Özel kural: Destek programı adı sonundaki "Programı" (veya Programi, PROGRAM vb.) kelimesini sil.
    if field_id == "destek_programi_adi":
        formatted = re.sub(r'\s+[pP][rR][oO][gG][rR][aA][mM][ıiIİ]?\s*$', '', formatted)
        
    return formatted


def _get_run_style_snapshot(run):
    """Bir run'ın font ve stil bilgisini hafif bir snapshot olarak döndürür."""
    if run is None:
        return None

    r_fonts = getattr(getattr(run._element, "rPr", None), "rFonts", None)
    return {
        "style": run.style,
        "font_name": run.font.name,
        "font_size": run.font.size,
        "ascii_font": r_fonts.get(qn("w:ascii")) if r_fonts is not None else None,
        "hansi_font": r_fonts.get(qn("w:hAnsi")) if r_fonts is not None else None,
        "eastasia_font": r_fonts.get(qn("w:eastAsia")) if r_fonts is not None else None,
        "cs_font": r_fonts.get(qn("w:cs")) if r_fonts is not None else None,
    }


def _get_reference_style_snapshot(runs, preferred_index=0):
    """Önce tercih edilen index'in gerisinden, sonra ilerisinden anlamlı run arar."""
    if not runs:
        return None

    preferred_index = max(0, min(preferred_index, len(runs) - 1))

    for idx in range(preferred_index, -1, -1):
        if runs[idx].text and runs[idx].text.strip():
            return _get_run_style_snapshot(runs[idx])

    for idx in range(preferred_index + 1, len(runs)):
        if runs[idx].text and runs[idx].text.strip():
            return _get_run_style_snapshot(runs[idx])

    for run in runs:
        if run.text:
            return _get_run_style_snapshot(run)

    return None


def _apply_style_snapshot(run, snapshot):
    """Snapshot'taki yazı tipi ve boyutunu yeni run'a uygular."""
    if not snapshot:
        return

    if snapshot.get("style") is not None:
        run.style = snapshot["style"]

    if snapshot.get("font_name"):
        run.font.name = snapshot["font_name"]
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        for key in ["ascii_font", "hansi_font", "eastasia_font", "cs_font"]:
            value = snapshot.get(key)
            if value:
                attr = {
                    "ascii_font": qn("w:ascii"),
                    "hansi_font": qn("w:hAnsi"),
                    "eastasia_font": qn("w:eastAsia"),
                    "cs_font": qn("w:cs"),
                }[key]
                r_fonts.set(attr, value)

    if snapshot.get("font_size") is not None:
        run.font.size = snapshot["font_size"]


def _add_run_with_style(para, text, snapshot=None, bold=None, italic=None):
    """Yeni run ekler ve mümkünse referans tipografisini korur."""
    run = para.add_run(text)
    _apply_style_snapshot(run, snapshot)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def _paragraph_has_numbering(para):
    """Paragraf Word numbering taşıyorsa True döndürür."""
    ppr = para._element.pPr
    return bool(ppr is not None and ppr.numPr is not None)


def _strip_leading_bullet_marker(text, marker):
    """Başta tekrarlayan a)/b) gibi markerları tekilleştirmek için temizler."""
    if not text:
        return ""
    marker = str(marker).strip()
    if len(marker) >= 2 and marker.endswith(")"):
        token = re.escape(marker[0]) + r'\s*\)'
    else:
        token = re.escape(marker)
    cleaned = re.sub(r'^\s*(?:' + token + r'\s*)+', '', str(text))
    return cleaned.strip()


def _set_bold_value_paragraph(para, value, prefix="", prefix_snapshot=None, value_snapshot=None):
    """Paragrafı opsiyonel prefix ve kalın değer ile yeniden yazar."""
    para.text = ""
    if prefix:
        _add_run_with_style(para, prefix, prefix_snapshot)

    lines = [line.strip() for line in str(value).splitlines() if line.strip()]
    if not lines:
        lines = [str(value).strip()]

    for idx, line in enumerate(lines):
        if idx > 0:
            _add_run_with_style(para, "", value_snapshot).add_break()
            if prefix:
                _add_run_with_style(para, prefix, prefix_snapshot)
        _add_run_with_style(para, line, value_snapshot, bold=True, italic=False)


def _rewrite_label_value_paragraph(para, value, label_text=None):
    """`Label: değer` biçimindeki satırı mevcut tipografiyle yeniden yazar."""
    existing_runs = list(para.runs)
    label_snapshot = _get_reference_style_snapshot(existing_runs, 0)
    value_snapshot = _get_reference_style_snapshot(existing_runs, len(existing_runs) - 1)

    if label_text is None:
        label_text = para.text.split(":", 1)[0] + ":"

    para.text = ""
    _add_run_with_style(para, label_text, label_snapshot)
    separator = "" if label_text.endswith((" ", "\t")) else " "
    _add_run_with_style(para, separator + str(value).strip(), value_snapshot, bold=True, italic=False)


def _clean_sozlesme_makami_adres_placeholder(text):
    """Şablonda kalan <Sözleşme Makamının Adresine> placeholder'ını temizler."""
    if not text:
        return text
    cleaned = re.sub(r'\s*<\s*Sözleşme Makamının Adresine\s*>\s*', ' ', str(text), flags=re.IGNORECASE)
    cleaned = re.sub(r'\s{2,}', ' ', cleaned)
    return cleaned


PRESERVE_YELLOW_TEXTS = {
    "Okudum, kabul ediyorum. .../.../20...",
    "İmza",
    "Teklif Veren (Yetkili imzası/imzaları ve kaşe)",
}

PRESERVE_YELLOW_SUBSTRINGS = {
    "Tedarikçinin/Hizmet Sunucusunun/Yapım Müteahhidinin Tam Resmi Adı",
    "Hukuki statüsü / ünvanı",
    "Resmi tescil numarası",
    "Açık resmi-tebligat adresi",
    "Vergi dairesi ve numarası",
}


def _ensure_page_break_before_heading(doc, heading_text):
    """Belirli bir başlığın her zaman yeni sayfadan başlamasını sağlar."""
    if not heading_text:
        return

    normalized_heading = heading_text.strip().lower()
    for para in doc.paragraphs:
        if para.text.strip().lower() == normalized_heading:
            para.paragraph_format.page_break_before = True
            return


def _ensure_page_break_before_paragraph_containing(doc, text_fragment):
    """İçeriğinde verilen metni barındıran ilk paragrafı yeni sayfadan başlatır."""
    if not text_fragment:
        return

    normalized_fragment = text_fragment.strip().lower()
    for para in doc.paragraphs:
        if normalized_fragment in para.text.strip().lower():
            para.paragraph_format.page_break_before = True
            return


def _move_davet_mektubu_section_break_before_teklif_dosyasi(doc):
    """
    İhaleye Davet Mektubu section break'ini, 'TEKLİF DOSYASI' başlığından hemen
    önceki paragrafa taşır. Böylece davet mektubu uzarsa sonraki section header'ı
    erken başlamaz.
    """
    in_davet = False
    source_para = None
    target_idx = None

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "İHALEYE DAVET MEKTUBU" in text:
            in_davet = True
        if in_davet and para._element.pPr is not None and para._element.pPr.sectPr is not None:
            source_para = para
            in_davet = False
        if text == "TEKLİF DOSYASI" and target_idx is None:
            target_idx = idx

    if source_para is None or target_idx is None or target_idx <= 0:
        return

    target_para = doc.paragraphs[target_idx - 1]
    source_ppr = source_para._element.get_or_add_pPr()
    source_sect = source_ppr.sectPr
    if source_sect is None:
        return

    target_ppr = target_para._element.get_or_add_pPr()
    if target_ppr.sectPr is not None:
        target_ppr.remove(target_ppr.sectPr)

    target_ppr.append(deepcopy(source_sect))
    source_ppr.remove(source_sect)


def generate_filled_document(template_path, form_data, output_dir="output"):
    """
    Şablon Word belgesini kullanıcı verileriyle doldurarak yeni belge oluşturur.
    """
    os.makedirs(output_dir, exist_ok=True)

    # Şablonu aç
    doc = Document(template_path)

    # İstatistikler
    stats = {
        "total_yellow": 0,
        "filled": 0,
        "instruction_skipped": 0,
        "unmatched": 0,
    }

    # Context bilgileri
    context = {
        "in_davet_mektubu": False,
        "fill_submission_address_next": False,
        "fill_submission_contact_next": False,
        "fill_info_address_next": False,
        "clear_info_address_extra_next": False,
        "fill_authority_person_next": False,
        "authority_signature_snapshot": None,
    }

    # --- PARAGRAFLAR ---
    paragraphs_to_remove = []
    
    for para in doc.paragraphs:
        # Davet mektubu kontrolü
        if "İHALEYE DAVET MEKTUBU" in para.text:
            context["in_davet_mektubu"] = True
            
        # Değerlendirme paragrafı a)/b) kontrolü
        if context.get("in_davet_mektubu"):
            text_norm = para.text.strip().lower()
            existing_runs = list(para.runs)

            if context.get("fill_submission_address_next"):
                kurum_adresi = _format_field_value("kurum_adresi", form_data.get("kurum_adresi", ""))
                context["fill_submission_address_next"] = False
                context["fill_submission_contact_next"] = True
                if kurum_adresi:
                    ref_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                    _set_bold_value_paragraph(
                        para,
                        kurum_adresi,
                        prefix="\t",
                        prefix_snapshot=ref_snapshot,
                        value_snapshot=ref_snapshot,
                    )
                    continue

            if context.get("fill_info_address_next"):
                kurum_adresi = _format_field_value("kurum_adresi", form_data.get("kurum_adresi", ""))
                context["fill_info_address_next"] = False
                context["clear_info_address_extra_next"] = True
                if kurum_adresi:
                    ref_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                    _set_bold_value_paragraph(
                        para,
                        kurum_adresi,
                        prefix="\t",
                        prefix_snapshot=ref_snapshot,
                        value_snapshot=ref_snapshot,
                    )
                    continue

            if context.get("clear_info_address_extra_next"):
                normalized = para.text.replace("\t", "").replace("_", "").strip()
                context["clear_info_address_extra_next"] = False
                if normalized == "":
                    para.text = ""
                    continue

            if context.get("fill_authority_person_next") and para.text.strip() == "":
                ilgili_personel = _format_field_value("kurum_ilgili_personel", form_data.get("kurum_ilgili_personel", ""))
                context["fill_authority_person_next"] = False
                ref_snapshot = context.get("authority_signature_snapshot")
                context["authority_signature_snapshot"] = None
                if ilgili_personel:
                    _set_bold_value_paragraph(
                        para,
                        ilgili_personel,
                        prefix="",
                        prefix_snapshot=ref_snapshot,
                        value_snapshot=ref_snapshot,
                    )
                context["in_davet_mektubu"] = False
                context["fill_submission_address_next"] = False
                context["fill_submission_contact_next"] = False
                context["fill_info_address_next"] = False
                context["clear_info_address_extra_next"] = False
                continue

            if context.get("fill_submission_contact_next") and "Telefon:" in para.text and "Faks:" in para.text:
                kurum_telefon = _format_field_value("kurum_telefon", form_data.get("kurum_telefon", ""))
                kurum_faks = _format_field_value("kurum_faks", form_data.get("kurum_faks", ""))
                phone_label_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                fax_label_snapshot = _get_reference_style_snapshot(existing_runs, len(existing_runs) - 1)
                para.text = ""
                _add_run_with_style(para, "\tTelefon: ", phone_label_snapshot)
                _add_run_with_style(para, kurum_telefon, phone_label_snapshot, bold=True, italic=False)
                _add_run_with_style(para, "\t\t\tFaks: ", fax_label_snapshot)
                _add_run_with_style(para, kurum_faks, fax_label_snapshot, bold=True, italic=False)
                context["fill_submission_contact_next"] = False
                continue

            if "İstenen formata uygun hazırlanmış teklifiniz aşağıdaki adrese gönderilmelidir:" in para.text:
                context["fill_submission_address_next"] = True

            if "Daha fazla bilgi aşağıdaki adresten elde edilebilir:" in para.text:
                context["fill_info_address_next"] = True

            if "Sözleşme Makamı Yetkilisi" in para.text:
                kurum_adi = _format_field_value("kurum_adi", form_data.get("kurum_adi", ""))
                if kurum_adi:
                    ref_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                    _set_bold_value_paragraph(
                        para,
                        kurum_adi,
                        prefix="",
                        prefix_snapshot=ref_snapshot,
                        value_snapshot=ref_snapshot,
                    )
                    context["fill_authority_person_next"] = True
                    context["authority_signature_snapshot"] = ref_snapshot
                    continue

            if "Proje Adı" in para.text and ":" in para.text and "_" in para.text:
                proje_adi = str(form_data.get("proje_adi", "")).strip()
                if proje_adi:
                    label_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                    value_snapshot = _get_reference_style_snapshot(existing_runs, len(existing_runs) - 1)
                    label = para.text.split(":", 1)[0] + ": "
                    para.text = ""
                    _add_run_with_style(para, label, label_snapshot)
                    _add_run_with_style(para, proje_adi, value_snapshot, bold=True, italic=False)
                    continue

            # Başlıktaki talimat metnini temizle (Parantez içi temizliği)
            # Daha sağlam bir regex ile tespiti yapalım (büyük/küçük harf ve Türkçe karakter toleranslı)
            if "DEĞERLENDİRME" in para.text and re.search(r'de[ğg]erlendirme.*(uygun|se[çc]iniz|siliniz)', para.text, flags=re.IGNORECASE):
                # 'DEĞERLENDİRME' kelimesini ve öncesini koru, sonrasını sil
                header_match = re.search(r'(.*DEĞERLENDİRME)', para.text, flags=re.IGNORECASE)
                if header_match:
                    header_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                    para.text = ""
                    _add_run_with_style(para, header_match.group(1).strip() + ": ", header_snapshot)
                continue # Bu paragraf için başka işlem yapma
            
            # Seçeneğe göre filtreleme
            turu = str(form_data.get("ihale_turu", ""))
            
            # İçerik bazlı tespit (markerlar auto-numbering nedeniyle p.text'te olmayabilir)
            is_a_content = "mal alımı" in text_norm and "en ucuz teklifi veren" in text_norm
            is_b_content = "hizmet alımlarında" in text_norm and "teknik değerlendirmenin" in text_norm
            
            if is_a_content:
                if turu == "Hizmet Alımı":
                    paragraphs_to_remove.append(para)
                else:
                    # a) şıkkını koru, prefix ekle/koru, gereksiz parantezi sil
                    line_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                    clean_text = _strip_leading_bullet_marker(para.text, "a)")
                    clean_text = clean_text.rstrip(")").strip()
                    para.text = ""
                    prefix = "" if _paragraph_has_numbering(para) else "a) "
                    _add_run_with_style(para, prefix + clean_text, line_snapshot)
                continue # Bu paragraf için başka işlem yapma
                
            elif is_b_content:
                if turu in ["Mal Alımı", "Yapım İşi"]:
                    paragraphs_to_remove.append(para)
                else:
                    # b) şıkkını koru, prefix ekle/koru, gereksiz parantezi sil
                    line_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                    clean_text = _strip_leading_bullet_marker(para.text, "b)")
                    clean_text = clean_text.rstrip(")").strip()
                    para.text = ""
                    prefix = "" if _paragraph_has_numbering(para) else "b) "
                    _add_run_with_style(para, prefix + clean_text, line_snapshot)
                continue # Bu paragraf için başka işlem yapma
                    
        # Madde 14: İSTEKLİLERE TALİMATLAR "Aşağıda yer alan maddeler..." silimi
        if "Aşağıda yer alan maddeler içerisindeki boş yerler" in para.text and "Diğer metinleri hiçbir şekilde değiştirmeyiniz" in para.text:
            paragraphs_to_remove.append(para)

        # Madde 1: Sözleşme Makamı bilgileri - placeholder uzunluğu çakışmalarını
        # etiket bazlı doldurma ile önle.
        if "a)  Adı/Ünvanı" in para.text:
            val = _format_field_value("kurum_adi", form_data.get("kurum_adi", ""))
            if val:
                _rewrite_label_value_paragraph(para, val, "a)  Adı/Ünvanı :")
                continue

        if "b)  Adresi:" in para.text:
            val = _format_field_value("kurum_adresi", form_data.get("kurum_adresi", ""))
            if val:
                _rewrite_label_value_paragraph(para, val, "b)  Adresi:")
                continue

        if "c)  Telefon numarası:" in para.text:
            val = _format_field_value("kurum_telefon", form_data.get("kurum_telefon", ""))
            if val:
                _rewrite_label_value_paragraph(para, val, "c)  Telefon numarası:")
                continue

        if "d)  Faks numarası:" in para.text:
            val = _format_field_value("kurum_faks", form_data.get("kurum_faks", ""))
            if val:
                _rewrite_label_value_paragraph(para, val, "d)  Faks numarası:")
                continue

        if "f)  İlgili personelinin adı-soyadı/unvanı:" in para.text:
            val = _format_field_value("kurum_ilgili_personel", form_data.get("kurum_ilgili_personel", ""))
            if val:
                _rewrite_label_value_paragraph(para, val, "f)  İlgili personelinin adı-soyadı/unvanı:")
                continue

        # Madde 2: İhale konusu işe ilişkin bilgiler - placeholder çakışmalarını
        # etiket bazlı doldurma ile önle.
        if "Projenin Adı:" in para.text:
            val = _format_field_value("proje_adi", form_data.get("proje_adi", ""))
            if val:
                _rewrite_label_value_paragraph(para, val, "Projenin Adı:")
                continue

        if "Fiziki Miktarı ve türü:" in para.text:
            val = _format_field_value("fiziki_miktar_ve_tur", form_data.get("fiziki_miktar_ve_tur", ""))
            if val:
                _rewrite_label_value_paragraph(para, val, "Fiziki Miktarı ve türü:")
                continue

        if "İşin/Teslimin Gerçekleştirileceği yer:" in para.text:
            val = _format_field_value("teslim_yeri", form_data.get("teslim_yeri", ""))
            if val:
                _rewrite_label_value_paragraph(para, val, "İşin/Teslimin Gerçekleştirileceği yer:")
                continue

        if "Alıma ait (varsa) diğer bilgiler:" in para.text:
            val = _format_field_value("diger_bilgiler", form_data.get("diger_bilgiler", ""))
            if val:
                _rewrite_label_value_paragraph(para, val, "Alıma ait (varsa) diğer bilgiler:")
                continue

        if "Bu Sözleşmenin Konusu" in para.text and "<Sözleşme Başlığı>" in para.text:
            existing_runs = list(para.runs)
            common_snapshot = _get_reference_style_snapshot(existing_runs, 0)
            yer = _format_field_value("is_yeri_il_ilce", form_data.get("is_yeri_il_ilce", ""))
            ihale_turu = _format_field_value("ihale_turu", form_data.get("ihale_turu", ""))
            ihale_konusu_val = form_data.get("ihale_konusu", "")
            if isinstance(ihale_konusu_val, list):
                konular = [str(v).strip() for v in ihale_konusu_val if str(v).strip()]
            else:
                konular = [line.strip() for line in str(ihale_konusu_val).splitlines() if line.strip()]
                if len(konular) <= 1 and "," in str(ihale_konusu_val):
                    konular = [part.strip() for part in str(ihale_konusu_val).split(",") if part.strip()]
            konu_metni = ", ".join(konular)
            if ihale_turu == "Yapım İşi":
                tur_eki = "'dir"
            else:
                tur_eki = "'dır"
            para.text = ""
            _add_run_with_style(
                para,
                f"Bu Sözleşmenin Konusu {yer}‘da uygulanacak {konu_metni} {ihale_turu}{tur_eki}. ",
                common_snapshot,
            )
            continue

        # Sözleşme kodu doldurma (Task 16)
        if "Sözleşme kodu:" in para.text and "<Ajans ile Yararlanıcı arasında" in para.text:
            val = str(form_data.get("sozlesme_kodu", ""))
            if val:
                existing_runs = list(para.runs)
                label_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                value_snapshot = _get_reference_style_snapshot(existing_runs, len(existing_runs) - 1)
                para.text = ""
                _add_run_with_style(para, "Sözleşme kodu: ", label_snapshot)
                _add_run_with_style(para, val, value_snapshot, bold=True, italic=False)
                continue # İşlem tamam

        # Madde 8: İhalenin yabancı isteklilere açıklığı
        if "İhalenin yabancı isteklilere açıklığı" in para.text or "İhaleye yabancı isteklilere açıklığı" in para.text:
            # Bir sonraki paragrafları kontrol et (Çünkü başlık ayrı paragraf olabilir)
            context["in_madde_8"] = True
            continue
            
        if context.get("in_madde_8") and ("yerli yabancı tüm isteklilere açıktır" in para.text or "sadece yerli isteklilere açıktır" in para.text):
            val = form_data.get("istekli_kapsamı")
            if val:
                existing_runs = list(para.runs)
                common_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                para.text = ""
                _add_run_with_style(para, str(val), common_snapshot, bold=True, italic=False)
            context["in_madde_8"] = False # İşlem tamam
            continue

        # Madde 18: Teklif ve sözleşme türü
        if "Teklif ve sözleşme türü" in para.text and "esaslı" in para.text:
            val = form_data.get("teklif_esasi")
            if val:
                # Talimatı temizle
                instruct = "(her lot için uygun olan yöntem seçilecek ve belirtilecektir)"
                for r in para.runs:
                    if instruct in r.text:
                        r.text = r.text.replace(instruct, "")
                
                # Esas metni ve "esaslı" ekini ayarla
                target = "götürü bedel / birim fiyat esaslı"
                if target in para.text:
                    for r in para.runs:
                        if target in r.text:
                            r.text = r.text.replace(target, f"{val} esaslı")
                            r.font.bold = True
                else:
                     # "götürü bedel / birim fiyat" ifadesini bul ve "esaslı" ekini koru/ekle
                     for r in para.runs:
                         if "götürü bedel / birim fiyat" in r.text:
                             r.text = r.text.replace("götürü bedel / birim fiyat", val)
                             r.font.bold = True
                             if "esaslı" not in r.text and "esaslı" not in para.text:
                                 r.text += " esaslı"
            continue
            
            
        # Madde 22: Teslim Yeri ve "Adresine" Kelimesi Eklenmesi
        if "Taahhütlü posta / kargo servisi" in para.text or "Ya da Sözleşme Makamına doğrudan elden" in para.text:
            teslim_yeri = form_data.get("teslim_yeri", "")
            if teslim_yeri:
                target_1 = "Taahhütlü posta / kargo servisi) ile"
                target_2 = "Ya da Sözleşme Makamına doğrudan elden"
                
                if target_1 in para.text:
                    # Para text'ini temizleyip yeni yapıyı kuralım
                    old_text = para.text
                    existing_runs = list(para.runs)
                    common_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                    para.text = ""
                    # "Taahhütlü posta / kargo servisi) ile" kısmına kadar olan yeri koru (eğer bir madde numarası varsa vb)
                    prefix = old_text.split(target_1)[0]
                    _add_run_with_style(para, prefix + target_1 + " ", common_snapshot)
                    _add_run_with_style(para, f"{teslim_yeri}", common_snapshot, bold=True, italic=False)
                    _add_run_with_style(para, " Adresine", common_snapshot)
                    # Geri kalan metni de ekle (eğer varsa)
                    suffix = _clean_sozlesme_makami_adres_placeholder(old_text.split(target_1)[-1])
                    if suffix and suffix != old_text:
                        _add_run_with_style(para, suffix, common_snapshot)
                    continue

                if target_2 in para.text:
                    old_text = para.text
                    existing_runs = list(para.runs)
                    common_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                    para.text = ""
                    prefix = old_text.split(target_2)[0]
                    _add_run_with_style(para, prefix + target_2 + " ", common_snapshot)
                    _add_run_with_style(para, f"{teslim_yeri}", common_snapshot, bold=True, italic=False)
                    _add_run_with_style(para, " Adresine", common_snapshot)
                    suffix = _clean_sozlesme_makami_adres_placeholder(old_text.split(target_2)[-1])
                    if suffix and suffix != old_text:
                        _add_run_with_style(para, suffix, common_snapshot)
                    continue

        # Madde 17: İhale Usulü Kontrolü a) ve b) maddeleri
        if "İhaleye davet mektubu" in para.text:
            usul = form_data.get("ihale_usulu")
            if usul == "Açık İhale Usulü":
                paragraphs_to_remove.append(para)
            else: # Pazarlık Usulü
                # Satırı tamamen temizleyip yeniden kurgulayalım
                existing_runs = list(para.runs)
                common_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                has_numbering = _paragraph_has_numbering(para)
                para.text = ""
                prefix = "" if has_numbering else "a) "
                _add_run_with_style(para, prefix + "İhaleye davet mektubu", common_snapshot)
            continue

        if "Teklif Dosyası" in para.text and "Sözleşme Taslağı" in para.text:
            usul = form_data.get("ihale_usulu")
            # İçeriği ayıkla (varsa b) veya a) prefixlerini sil)
            clean_content = para.text.strip()
            clean_content = _strip_leading_bullet_marker(clean_content, "a)")
            clean_content = _strip_leading_bullet_marker(clean_content, "b)")
            
            existing_runs = list(para.runs)
            common_snapshot = _get_reference_style_snapshot(existing_runs, 0)
            has_numbering = _paragraph_has_numbering(para)
            para.text = ""
            if usul == "Açık İhale Usulü":
                prefix = "" if has_numbering else "a) "
                _add_run_with_style(para, prefix + clean_content, common_snapshot)
            else: # Pazarlık Usulü
                prefix = "" if has_numbering else "b) "
                _add_run_with_style(para, prefix + clean_content, common_snapshot)
            continue

        # Madde 17 f) bendi: Geçici Teminat Seçimi
        if "f) İstenmesi halinde Md. 26" in para.text and "geçici teminat" in para.text:
            val = form_data.get("gecici_teminat")
            if val:
                existing_runs = list(para.runs)
                label_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                value_snapshot = _get_reference_style_snapshot(existing_runs, len(existing_runs) - 1)
                para.text = ""
                _add_run_with_style(para, "f)\tİstenmesi halinde Md. 26’daki koşullara uygun sunulmuş geçici teminat ", label_snapshot)
                _add_run_with_style(para, f"{val}.", value_snapshot, bold=True, italic=False)
                continue

        # Madde 17 h) bendi: Kesin Teminat Seçimi
        if "h) İstenmesi halinde Genel Koşullar md. 29" in para.text and "kesin teminat" in para.text:
            val = form_data.get("kesin_teminat")
            if val:
                existing_runs = list(para.runs)
                label_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                value_snapshot = _get_reference_style_snapshot(existing_runs, len(existing_runs) - 1)
                para.text = ""
                _add_run_with_style(para, "h)\tİstenmesi halinde Genel Koşullar md. 29’da ve Taslak Sözleşme Özel Koşullar md. 8.1’de şartları tanımlanmış kesin teminat ", label_snapshot)
                _add_run_with_style(para, f"{val}.", value_snapshot, bold=True, italic=False)
                continue

        # Madde 25: Sözleşme ve Özel Koşullar - Kesin Teminat Bölümü (Task 25)
        if "Yüklenici tarafından sözleşme imzalama aşamasında kesin teminat" in para.text:
            val = form_data.get("kesin_teminat")
            if val:
                # "Kesin teminat ve Sigorta" başlığı aynı paragrafta olabilir, koruyalım
                prefix = ""
                if "Kesin teminat ve Sigorta" in para.text:
                    prefix = "Kesin teminat ve Sigorta "
                
                existing_runs = list(para.runs)
                common_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                para.text = ""
                _add_run_with_style(para, f"{prefix}Bu sözleşme kapsamında işin ihale edildiği Yüklenici tarafından sözleşme imzalama aşamasında kesin teminat ", common_snapshot)
                _add_run_with_style(para, f"{val}", common_snapshot, bold=True, italic=False)
                _add_run_with_style(para, ".", common_snapshot)
                continue
                
        # Madde 24: Ön Ödeme Paragrafı Yeniden Yazımı (Task 24)
        if "Sözleşme kapsamında ön ödeme" in para.text:
            val = form_data.get("on_odeme")
            existing_runs = list(para.runs)
            common_snapshot = _get_reference_style_snapshot(existing_runs, 0)
            if val == "Yapılmayacaktır":
                para.text = ""
                _add_run_with_style(para, "Sözleşme kapsamında ön ödeme ", common_snapshot)
                _add_run_with_style(para, "yapılmayacaktır", common_snapshot, bold=True, italic=False)
                _add_run_with_style(para, ".", common_snapshot)
            elif val == "Yapılacaktır":
                oran = form_data.get("on_odeme_orani", "")
                para.text = ""
                _add_run_with_style(para, "Sözleşme kapsamında ön ödeme ", common_snapshot)
                _add_run_with_style(para, "yapılacaktır", common_snapshot, bold=True, italic=False)
                _add_run_with_style(para, ". Ön ödeme miktarı sözleşme bedelinin % ", common_snapshot)
                _add_run_with_style(para, f"{oran}", common_snapshot, bold=True, italic=False)
                _add_run_with_style(para, "’sı olan ……………….. TL’dir. Ön ödeme, sözleşme imza tarihinden sonra 15 gün içerisinde belirlenen ön ödeme tutarı kadar avans teminat mektubunun sunulmasını takiben yapılacaktır.", common_snapshot)
            continue
                        
        # Madde 25: Kesin Teminat Oranı İkinci Cümle
        if "Kesin teminat tutarı sözleşme bedelinin" in para.text and "olmalıdır" in para.text:
            val = form_data.get("kesin_teminat")
            if val == "İSTENMEMEKTEDİR":
                paragraphs_to_remove.append(para)
                continue
            elif val == "İSTENMEKTEDİR":
                oran = form_data.get("kesin_teminat_orani", "")
                existing_runs = list(para.runs)
                common_snapshot = _get_reference_style_snapshot(existing_runs, 0)
                para.text = ""
                _add_run_with_style(para, "Kesin teminat tutarı sözleşme bedelinin % ", common_snapshot)
                _add_run_with_style(para, f"{oran}", common_snapshot, bold=True, italic=False)
                _add_run_with_style(para, "’sı kadar olmalıdır.", common_snapshot)
                continue

        # Madde 34, 36, 37: İdari Uygunluk / Sözleşme Adı / Teknik Tablo birleşik alan (< Proje adı >Projesi için...)
        p_text_lower = para.text.lower()
        if "proje ad" in p_text_lower and "projesi i" in p_text_lower and ("ad" in p_text_lower or "sözleşme" in p_text_lower):
            p_adi = str(form_data.get("proje_adi", ""))
            i_turu = str(form_data.get("ihale_turu", ""))
            p_adi_clean = re.sub(r'(?i)\s+projesi$', '', p_adi)
            
            # Label ve sekmeleri koru
            if "Sözleşme adı:" in para.text:
                label = "Sözleşme adı: "
            elif "Sözleşme başlı" in para.text:
                # Şablonda 'Sözleşme başlı	:' şeklinde (tab ile) geçiyor
                label = "Sözleşme başlı\t: "
            else:
                # 'Adı:' kısmında template'te \t\t var
                label = "Adı:\t\t"
            
            existing_runs = list(para.runs)
            label_snapshot = _get_reference_style_snapshot(existing_runs, 0)
            value_snapshot = _get_reference_style_snapshot(existing_runs, len(existing_runs) - 1)
            para.text = ""
            _add_run_with_style(para, label, label_snapshot)
            _add_run_with_style(para, f"{p_adi_clean} Projesi için {i_turu}", value_snapshot, bold=True, italic=False)
            continue
            
        _process_paragraph_runs(para, form_data, stats, context)

    # Remove the explicitly deleted paragraphs correctly
    for p in paragraphs_to_remove:
        try:
            p1 = p._element
            p1.getparent().remove(p1)
            p._p = p._element = None
        except Exception:
            pass

    # --- TABLOLAR ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph_runs(para, form_data, stats, context)

    # --- HEADER/FOOTER ---
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header and header.is_linked_to_previous is False:
                for para in header.paragraphs:
                    _process_paragraph_runs(para, form_data, stats, context)
        for footer in [section.footer, section.first_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for para in footer.paragraphs:
                    _process_paragraph_runs(para, form_data, stats, context)

    # Davet mektubu ile sonraki section arasındaki gerçek section break'i,
    # 'TEKLİF DOSYASI' başlığından hemen önceye taşı.
    _move_davet_mektubu_section_break_before_teklif_dosyasi(doc)

    # Dosya adı oluştur
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    kurum_adi = form_data.get("kurum_adi", "ihale")
    safe_name = re.sub(r'[^\w\s-]', '', kurum_adi).strip().replace(' ', '_')[:50]
    docx_filename = f"İhale_Dosyası_{safe_name}_{timestamp}.docx"
    docx_path = os.path.join(output_dir, docx_filename)

    # Kaydet
    doc.save(docx_path)

    result = {
        "docx": os.path.abspath(docx_path),
        "stats": stats,
    }

    # PDF dönüşümü dene
    try:
        pdf_path = _convert_to_pdf(docx_path)
        if pdf_path:
            result["pdf"] = pdf_path
    except Exception as e:
        result["pdf_error"] = str(e)

    return result


def _process_paragraph_runs(para, form_data, stats, context):
    runs = list(para.runs)
    if not runs:
        return

    para_text_normalized = _normalize_text(para.text)
    if para_text_normalized in PRESERVE_YELLOW_TEXTS:
        for r in runs:
            r.font.highlight_color = None
        return

    if any(token in para.text for token in PRESERVE_YELLOW_SUBSTRINGS):
        for r in runs:
            if r.font.highlight_color == 7:
                r.font.highlight_color = None
        return

    # Önce "(i)", "(ii)", "(iii)" durumunu kontrol edelim (Davet mektubu mal listesi)
    if context.get("in_davet_mektubu") and any(m in para.text for m in ["(i)", "(ii)", "(iii)"]):
        # Eğer paragrafta sadece (i) ____________________ vb varsa
        ihale_konusu_val = form_data.get("ihale_konusu", "")
        if isinstance(ihale_konusu_val, list):
            lines = [v.strip() for v in ihale_konusu_val if v and v.strip()]
        else:
            lines = [line.strip() for line in str(ihale_konusu_val).splitlines() if line.strip()]
        
        for idx, marker in enumerate(["(i)", "(ii)", "(iii)"]):
            if marker in para.text:
                if idx < len(lines):
                    # Replace _________ with the line item
                    for r in runs:
                        r.text = r.text.replace("____________________", lines[idx])
                        r.text = r.text.replace("_________________", lines[idx])
                        r.font.bold = True
                        r.font.italic = False
                else:
                    # Clear it completely if no item
                    for r in runs:
                        if marker in r.text or "_" in r.text:
                            r.text = ""

    # Madde 15: "e) Elektronik posta adresi…" -> "e) Elektronik posta adresi: info@kurum.com"
    if ("e) " in para.text or "e)" in para.text) and "Elektronik posta adresi" in para.text:
        email_val = str(form_data.get("kurum_eposta", ""))
        if email_val:
            existing_runs = list(para.runs)
            label_snapshot = _get_reference_style_snapshot(existing_runs, 0)
            value_snapshot = _get_reference_style_snapshot(existing_runs, len(existing_runs) - 1)
            # Mevcut metni temizle (prefix hariç veya komple)
            # Daha güvenli olması için paragrafı yeniden kurgulayalım
            para.text = ""
            _add_run_with_style(para, "e) Elektronik posta adresi: ", label_snapshot)
            _add_run_with_style(para, email_val, value_snapshot, bold=True, italic=False)
            return # Bu paragraf işlendi
        else:
            # Email yoksa sadece noktalamayı düzelt
            for r in runs:
                r.text = r.text.replace("adresi…", "adresi:").replace("...", ":").replace("…", ":")

    # Yayın Referansı Alanlarını Doldurma (Merkezi Kontrol: Task 27, 30, 32, 48, 34)
    # Farklı casing varyasyonları için (Yayın Referansı, YAyın referansı vb.) normalize kontrol
    para_norm = para.text.lower().replace('ı', 'i').replace('i̇', 'i')
    has_ref_label = "yayin referans" in para_norm
    has_placeholder = any(x in para_norm or x in para.text.lower() for x in ["<sözleşme no", "<sozlesme no", "<ihale no"])
    
    if has_ref_label and has_placeholder:
        sz_val = str(form_data.get("sozlesme_kodu", ""))
        if sz_val:
            existing_runs = list(para.runs)
            label_snapshot = _get_reference_style_snapshot(existing_runs, 0)
            value_snapshot = _get_reference_style_snapshot(existing_runs, len(existing_runs) - 1)
            para.text = ""
            _add_run_with_style(para, "Yayın Referansı: ", label_snapshot)
            _add_run_with_style(para, sz_val, value_snapshot, bold=True, italic=False)
            return # İşlem tamam

    # Teklif Sunum Formu - Taahhütname İhale Türü Uyarlaması (Task 31)
    if "imza atmaya yetkili kişisi olarak" in para.text:
        i_turu = str(form_data.get("ihale_turu", ""))
        phrase = ""
        if "Mal Alımı" in i_turu:
            phrase = "malları tedarik etmeyi"
        elif "Hizmet Alımı" in i_turu:
            phrase = "hizmetleri sağlamayı"
        elif "Yapım İşi" in i_turu:
            phrase = "yapım işini üstlenmeyi"
        
        if phrase:
            full_txt = para.text
            start_mark = "<hizmetleri"
            if start_mark in full_txt:
                start_idx = full_txt.find(start_mark)
                end_idx = full_txt.find(">", start_idx)
                if start_idx != -1 and end_idx != -1:
                    actual_placeholder = full_txt[start_idx:end_idx+1]
                    _replace_text_across_runs(para, actual_placeholder, phrase)
            
    # Madde 34: ">Projesi için" boşluk düzenlemesi
    for r in runs:
        if "Projesi için" in r.text:
            r.text = r.text.replace(">Projesi için", "> Projesi için").replace("Projesi için", " Projesi için").replace("  Projesi için", " Projesi için")

    # Madde 16, 18, 19, 20, 30, 38, 40: Gereksiz talimat metinlerini temizle
    for old_txt in [
        "<Örnek: TR21-11-İKT-1-XX>", "Örnek: TR21-11-İKT-1-XX",
        "<Örnek: TR21-23-İKT-XX/01>", "Örnek: TR21-23-İKT-XX/01",
        "<ÖRNEK: TR21-23-İKT-XX/01>", "ÖRNEK: TR21-23-İKT-XX/01",
        "< ÖRNEK: TR21-23-İKT-XX/01 >", "<ÖRNEK: TR21-23-İKT-XX/01 >",
        "<uygun olan seçeneği seçiniz >", "<uygun olan seçeneği seçiniz>",
        "uygun olan seçeneği seçiniz",
        "(Uygun olanı seçiniz)>", "(Uygun olanı seçiniz)", "Uygun olanı seçiniz",
        "(Yararlanıcı ihale konusu"
    ]:
        if old_txt in para.text:
            if old_txt == "(Yararlanıcı ihale konusu":
                 # Tam parantez içini temizle
                 para.text = re.sub(r'\(Yararlanıcı ihale konusu.*?\)', '', para.text)
            else:
                _replace_text_across_runs(para, old_txt, "")

    i = 0
    while i < len(runs):
        run = runs[i]

        # Sarı vurgulu mu?
        if run.font.highlight_color is not None and run.font.highlight_color == 7:
            stats["total_yellow"] += 1

            # Ardışık sarı run'ları birleştir
            combined_text = run.text
            j = i + 1
            yellow_runs = [run]

            while j < len(runs):
                next_run = runs[j]
                if (next_run.font.highlight_color is not None and
                        next_run.font.highlight_color == 7):
                    combined_text += next_run.text
                    yellow_runs.append(next_run)
                    stats["total_yellow"] += 1
                    j += 1
                else:
                    break

            # Eşleştir
            field_id = _match_yellow_to_field(combined_text)

            if (
                context.get("in_davet_mektubu")
                and "İhale konusu işin niteliğine göre istenen bilgi" in combined_text
            ):
                for r_item in yellow_runs:
                    r_item.text = ""
                    r_item.font.highlight_color = None
                if i - 1 >= 0 and runs[i - 1].text.isspace():
                    runs[i - 1].text = ""
                if i - 2 >= 0 and runs[i - 2].text.strip() == ".":
                    runs[i - 2].text = ""
                if j < len(runs) and runs[j].text.isspace():
                    runs[j].text = ""
                stats["instruction_skipped"] += 1
                i = j
                continue

            # Özel context geçersiz kılmaları:
            if context.get("in_davet_mektubu") and field_id == "ihale_tarihi":
                # Davet mektubundaki tarih alanı genellikle …./…./20… formatındadır
                if re.search(r'\d{4}|…', combined_text):
                    field_id = "davet_tarihi"

            if field_id is not None:
                if "Sözleşme Makamının (Mali Destek Yararlanıcısının) resmi adı ve adresi" in combined_text:
                    kurum_adi = _format_field_value("kurum_adi", form_data.get("kurum_adi", ""))
                    kurum_adresi = _format_field_value("kurum_adresi", form_data.get("kurum_adresi", ""))
                    ref_snapshot = _get_reference_style_snapshot(runs, i)
                    para.text = ""
                    if kurum_adi:
                        _add_run_with_style(para, kurum_adi, ref_snapshot, bold=True, italic=False)
                    if kurum_adresi:
                        if kurum_adi:
                            _add_run_with_style(para, "", ref_snapshot).add_break()
                        _add_run_with_style(para, kurum_adresi, ref_snapshot, bold=True, italic=False)
                    stats["filled"] += 1
                    i = j
                    continue

                # Madde 34: Birleşik alan kontrolü (< Proje adı >Projesi için <Mal Alımı...>)
                if "Proje adı" in combined_text and "Projesi için" in combined_text:
                    p_adi = str(form_data.get("proje_adi", ""))
                    i_turu = str(form_data.get("ihale_turu", ""))
                    p_adi_clean = re.sub(r'(?i)\s+projesi$', '', p_adi)
                    value = f"{p_adi_clean} Projesi için {i_turu}"
                else:
                    value = form_data.get(field_id, "")
                
                # Özel Biçimlendirmeler:
                
                # İhale Konusu başka bir yerdeyse (tek satırda) virgülle birleştir
                if field_id == "ihale_konusu":
                    if isinstance(value, list):
                        value = ", ".join([str(v).strip() for v in value if str(v).strip()])
                    elif "\n" in str(value):
                        lines = [line.strip() for line in str(value).splitlines() if line.strip()]
                        value = ", ".join(lines)
                    
                # Yeterlik Değerlendirme (Davet mektubu teknik teklif paragrafında)
                if field_id == "yeterlik_degerlendirme":
                    if "\n" in str(value):
                        lines = [line.strip() for line in str(value).splitlines() if line.strip()]
                        value = ", ".join(lines)
                    if not str(value).endswith("yeterlik değerlendirmesinde kullanılacaktır."):
                        value = str(value).rstrip(".") + " yeterlik değerlendirmesinde kullanılacaktır."
                
                # Saat ise "Teklif teslimi..." gibi paragrafta isek tarih ile saati birleştirebiliriz
                if field_id == "ihale_tarihi":
                    text_lower = para.text.lower()
                    if "tarih" in text_lower and "saat" in text_lower:
                        ihale_saati = _format_field_value("ihale_saati", form_data.get("ihale_saati", ""))
                        if ihale_saati and ihale_saati not in str(value):
                            value = str(value) + " " + ihale_saati
                        
                # Madde 21: Teklif Esası -> "esaslı" ekle
                if field_id == "teklif_esasi":
                    value = str(value) + " esaslı"
                    
                # Madde 22: Taahhütlü Posta / Elden Teslim -> " Adresine" ekle
                if field_id == "teslim_yeri":
                    if "Taahhütlü posta" in para.text or "doğrudan elden" in para.text:
                        value = str(value) + " Adresine"
                        
                # Madde 27, 33: "İsteklinin adı" dots should remain dots (Signature lines)
                lower_para = para.text.lower()
                if "isteklinin" in lower_para and "ad" in lower_para:
                    stats["instruction_skipped"] += 1
                    yellow_runs[0].font.highlight_color = None
                    for extra_run in yellow_runs[1:]:
                        extra_run.font.highlight_color = None
                    i = j
                    continue

                # Madde 35: Lot Numarası alanını tamamen sil
                norm_combined = combined_text.lower().strip()
                if "lot numarası" in norm_combined or "lot numarasi" in norm_combined:
                    for r_item in yellow_runs:
                        r_item.text = ""
                        r_item.font.highlight_color = None
                    i = j
                    continue

                # Madde 23: Sözleşme başlığında İhale Türü BÜYÜK ve BOLD olmalı
                if field_id == "ihale_turu" and "SÖZLEŞMESİ" in para.text:
                    value = str(value).upper()
                    
                # Madde 26, 29, 31, 34, 36, 37: Proje adında 'Projesi' tekrarını önle
                if field_id == "proje_adi":
                    value = re.sub(r'(?i)\s+projesi$', '', str(value))
                    
                # Madde 39: Taahhütname cümlesi ihale türüne göre özel cümle ataması
                if field_id == "ihale_turu" and ("hizmetleri sağlamayı" in combined_text or "malları tedarik etmeyi" in combined_text):
                    if str(value) == "Mal Alımı": value = "malları tedarik etmeyi"
                    elif str(value) == "Hizmet Alımı": value = "hizmetleri sağlamayı"
                    elif str(value) == "Yapım İşi": value = "yapım işini üstlenmeyi"
                
                formatted_value = _format_field_value(field_id, value)

                if formatted_value:
                    # İlk run'a değeri yaz
                    yellow_runs[0].text = formatted_value
                    yellow_runs[0].font.highlight_color = None  # Sarıyı kaldır
                    
                    # Bold olmasını ve İtalik olmamasını sağla
                    yellow_runs[0].font.bold = True
                    yellow_runs[0].font.italic = False

                    # Geri kalan run'ları temizle
                    for extra_run in yellow_runs[1:]:
                        extra_run.text = ""
                        extra_run.font.highlight_color = None

                    # --- ETRAFIAKİ < > İŞARETLERİNİ VE EKLERİ TEMİZLE/DÜZELT ---
                    
                    # Geriye dönük < sil
                    for prev_idx in range(i-1, -1, -1):
                        txt = runs[prev_idx].text
                        if '<' in txt:
                            runs[prev_idx].text = txt.replace('<', '')
                            break
                        if len(txt.strip()) > 0:
                            break

                    # İleriye dönük > sil ve son ekleri (suffix) akıllıca yakala
                    for next_idx in range(j, len(runs)):
                        txt = runs[next_idx].text
                        if not txt: continue
                        
                        original_was_just_marker = False
                        if '>' in txt:
                            if txt.strip() == '>':
                                original_was_just_marker = True
                            txt = txt.replace('>', '')
                            runs[next_idx].text = txt
                            
                        # Eğer bu run hala metin içeriyorsa veya sadece '>' değil idiyse (yani peşinden bir şey gelebilir)
                        # Suffix tespiti için sonraki birkaç run'ı birleştirerek kontrol et (bölünmüş run desteği)
                        if txt.strip() or original_was_just_marker:
                            combined_suffix = ""
                            suffix_runs = []
                            # Sonraki 5 run'ı birleştir (genelde suffix 2-3 run'a yayılır: ' | n | e)
                            for k in range(next_idx, min(next_idx + 5, len(runs))):
                                combined_suffix += runs[k].text
                                suffix_runs.append(runs[k])
                                if " " in runs[k].text: break # Boşluk gelirse kelime bitmiştir
                                
                            # Regex: Tırnak ile başlayan (opsiyonel boşluktan sonra) ve harflerle devam eden yapı
                            # (not: şablonda hem düz ('), hem kıvrık sağ (’) hem de kıvrık sol (‘) tırnaklar olabiliyor!)
                            m = re.match(r'^(\s*)([\'’‘][a-zçğıöşü]+)(.*)', combined_suffix, flags=re.IGNORECASE)
                            if m:
                                target_suffix = None
                                if field_id == "is_yeri_il_ilce":
                                    target_suffix = get_locative_suffix(formatted_value)
                                elif field_id == "ihale_saati":
                                    target_suffix = get_dative_suffix(formatted_value)
                                
                                if target_suffix:
                                    quote = m.group(2)[0]
                                    # Yeni eki oluştur (varsa baştaki boşluğu silerek yapıştırır)
                                    full_new = quote + target_suffix[1:] + m.group(3)
                                    suffix_runs[0].text = full_new
                                    # Diğer birleştirilen run'ları temizle
                                    for sr in suffix_runs[1:]:
                                        sr.text = ""
                                    break # Suffix işlendi
                            
                            if txt.strip(): # Eğer suffix değilse ama metin varsa aramayı durdur
                                break

                    stats["filled"] += 1
                else:
                    stats["unmatched"] += 1
            else:
                is_instruction = False
                normalized = _normalize_text(combined_text)
                for instruction in INSTRUCTION_FIELDS:
                    if instruction in normalized or normalized in instruction:
                        is_instruction = True
                        break

                if is_instruction:
                    for yr in yellow_runs:
                        if yr.text.strip().startswith("___"):
                            yr.font.highlight_color = None
                        else:
                            yr.text = ""
                            yr.font.highlight_color = None
                        
                    # Talimat çevresindeki < ve > işaretlerini temizle
                    for prev_idx in range(i-1, -1, -1):
                        txt = runs[prev_idx].text
                        if '<' in txt:
                            runs[prev_idx].text = txt.replace('<', '')
                            break
                        if len(txt.strip()) > 0:
                            break
                            
                    for next_idx in range(j, len(runs)):
                        txt = runs[next_idx].text
                        if '>' in txt:
                            runs[next_idx].text = txt.replace('>', '')
                            break
                        if len(txt.strip()) > 0:
                            break
                    
                    stats["instruction_skipped"] += 1
                else:
                    stats["unmatched"] += 1

            i = j
        else:
            i += 1


def _convert_to_pdf(docx_path):
    import subprocess
    import platform

    abs_docx = os.path.abspath(docx_path)
    output_dir = os.path.dirname(abs_docx)
    pdf_path = abs_docx.replace(".docx", ".pdf")

    try:
        libreoffice_cmd = "libreoffice"
        if platform.system() == "Windows":
            lo_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
            for lo_path in lo_paths:
                if os.path.exists(lo_path):
                    libreoffice_cmd = lo_path
                    break

        result = subprocess.run(
            [
                libreoffice_cmd,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                abs_docx,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if os.path.exists(pdf_path):
            return pdf_path
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    except Exception:
        pass

    if platform.system() == "Windows":
        try:
            import win32com.client

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(abs_docx)
            doc.SaveAs2(pdf_path, FileFormat=17)
            doc.Close()
            word.Quit()
            return pdf_path
        except ImportError:
            pass
        except Exception:
            try:
                word.Quit()
            except Exception:
                pass

    return None


def validate_form_data(form_data):
    errors = []

    for field_id, field_info in UNIQUE_FIELDS.items():
        if field_info.get("required", False):
            value = form_data.get(field_id)
            if value is None or value == "" or value == []:
                errors.append(
                    f"'{field_info['label']}' alanı zorunludur."
                )

        if field_info.get("type") == "number" and field_id in form_data:
            value = form_data[field_id]
            if value is not None and value != "":
                try:
                    num_val = float(value)
                    min_val = field_info.get("min_value")
                    max_val = field_info.get("max_value")
                    if min_val is not None and num_val < min_val:
                        errors.append(
                            f"'{field_info['label']}' en az {min_val} olmalıdır."
                        )
                    if max_val is not None and num_val > max_val:
                        errors.append(
                            f"'{field_info['label']}' en fazla {max_val} olmalıdır."
                        )
                except (ValueError, TypeError):
                    errors.append(
                        f"'{field_info['label']}' geçerli bir sayı olmalıdır."
                    )

    return len(errors) == 0, errors


def _replace_text_across_runs(paragraph, old_text, new_text):
    """
    Paragraf içindeki metni (farklı run'lara bölünmüş olsa dahi) bulur ve değiştirir.
    """
    if not old_text or old_text not in paragraph.text:
        return

    # Önce tekil run'larda dene (formatı korumak için)
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            if old_text not in paragraph.text:
                return

    # Eğer hala metin varsa, paragraf düzeyinde değiştir. 
    # Not: Bu yöntem paragraftaki run-level formatları (bold/italic vb.) 
    # eğer metin birden fazla run'a yayılmışsa sıfırlayabilir.
    paragraph.text = paragraph.text.replace(old_text, new_text)
