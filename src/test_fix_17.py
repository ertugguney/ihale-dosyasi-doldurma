import os
import sys
import csv
from docx import Document
import re

# src klasörünü path'e ekle
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), 'src')))

from doc_generator import generate_filled_document

def load_test_data(csv_path):
    data = {}
    with open(csv_path, mode='r', encoding='utf-8-sig') as f:
        reader = csv.DictReader(f, delimiter=';')
        for row in reader:
            data[row['Alan ID']] = row['Değer']
    return data

def check_madde_17(docx_path, usul):
    doc = Document(docx_path)
    found_header = False
    items = []
    
    # Sadece "İhale dosyası aşağıdaki belgelerden oluşmaktadır" başlığı altındaki 2 satıra bak
    for i, para in enumerate(doc.paragraphs):
        if "İhale dosyası aşağıdaki belgelerden oluşmaktadır" in para.text:
            found_header = True
            # Sonraki 2 paragrafı kontrol et
            for j in range(i+1, i+4): # Biraz pay bırakalım
                p_text = doc.paragraphs[j].text.strip()
                if p_text:
                    items.append(p_text)
                    if len(items) >= (2 if usul == "Pazarlık Usulü" else 1):
                        break
            break

    print(f"\n--- {usul} için Madde 17 Kontrolü ---")
    print(f"Bulunan maddeler: {items}")
    
    success = True
    if usul == "Pazarlık Usulü":
        # a) İhaleye davet mektubu (talimatsız)
        # b) Teklif Dosyası (...)
        if len(items) < 2:
            print("HATA: En az 2 madde bulunmalıydı!")
            success = False
        else:
            # Etiketi (a) ) ayıklayıp içeriğe bak
            content_a = items[0][2:].strip()
            if "İhaleye davet mektubu" not in content_a or "(" in content_a or ")" in content_a:
                print(f"HATA: 1. madde hatalı: '{items[0]}'")
                success = False
            if not items[0].startswith("a)"):
                print(f"HATA: 1. madde a) ile başlamalı: '{items[0]}'")
                success = False
            if "Teklif Dosyası" not in items[1] or not items[1].startswith("b)"):
                print(f"HATA: 2. madde b) Teklif Dosyası olmalı: '{items[1]}'")
                success = False
    else: # Açık İhale Usulü
        # Sadece a) Teklif Dosyası (...) olmalı (davet mektubu silinmiş olmalı)
        if len(items) == 0:
            print("HATA: Hiç madde bulunamadı!")
            success = False
        else:
            if "Teklif Dosyası" not in items[0]:
                print(f"HATA: 1. madde Teklif Dosyası olmalı: '{items[0]}'")
                success = False
            if not items[0].startswith("a)"):
                print(f"HATA: 1. madde a) ile başlamalı (re-indexing): '{items[0]}'")
                success = False
            if any("İhaleye davet mektubu" in item for item in items):
                print("HATA: İhaleye davet mektubu silinmemiş!")
                success = False

    return success

if __name__ == "__main__":
    csv_path = r"C:\Users\eguney\Desktop\ihale\output\ihale_form_verileri_20260225.csv"
    template_path = r"Taslak İhale Dosyası.docx"
    output_dir = r"C:\Users\eguney\Desktop\ihale\output\test_fix_17"
    
    if not os.path.exists(csv_path):
        print(f"HATA: CSV bulunamadı: {csv_path}")
        sys.exit(1)
        
    base_data = load_test_data(csv_path)
    
    # 1. TEST: Pazarlık Usulü
    data_pazarlik = base_data.copy()
    data_pazarlik["ihale_usulu"] = "Pazarlık Usulü"
    res_pazarlik = generate_filled_document(template_path, data_pazarlik, output_dir=os.path.join(output_dir, "pazarlik"))
    s1 = check_madde_17(res_pazarlik['docx'], "Pazarlık Usulü")
    
    # 2. TEST: Açık İhale Usulü
    data_acik = base_data.copy()
    data_acik["ihale_usulu"] = "Açık İhale Usulü"
    res_acik = generate_filled_document(template_path, data_acik, output_dir=os.path.join(output_dir, "acik"))
    s2 = check_madde_17(res_acik['docx'], "Açık İhale Usulü")
    
    if s1 and s2:
        print("\n=== TÜM TESTLER BAŞARILI ===")
    else:
        print("\n=== TESTLERDE HATALAR VAR ===")
        sys.exit(1)
